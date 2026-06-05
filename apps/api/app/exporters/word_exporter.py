"""DOCX export functionality."""

import os
from datetime import datetime, timezone
from typing import Optional

from app.core.config import settings


class WordExporter:
    """Export content to Microsoft Word (.docx) format."""

    async def export(
        self,
        content: str,
        filename: str,
        title: Optional[str] = None,
        author: str = "3D Wall Platform",
    ) -> str:
        """Export markdown-like content to a Word document.

        Returns the file path of the generated document.
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()

        # Set default style
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x20, 0x21, 0x24)

        # Add document properties
        doc.core_properties.title = title or filename
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now(timezone.utc)

        # Process content line by line
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.startswith("# "):
                h = doc.add_heading(stripped[2:], level=1)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            elif stripped.startswith("## "):
                h = doc.add_heading(stripped[3:], level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("#### "):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith(tuple(f"{i}. " for i in range(1, 10))):
                doc.add_paragraph(stripped, style="List Number")
            elif stripped.startswith("---"):
                doc.add_paragraph("")
                p = doc.add_paragraph()
                run = p.add_run("—" * 40)
                run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            elif stripped.startswith("| "):
                # Table row — render as formatted text for simplicity
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            else:
                # Handle bold/italic inline
                p = doc.add_paragraph()
                _add_formatted_text(p, stripped)

        # Add footer
        doc.add_paragraph("")
        footer = doc.add_paragraph(
            f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
            f" by {author}"
        )
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Save to storage
        storage_dir = os.path.abspath(settings.storage_path)
        os.makedirs(storage_dir, exist_ok=True)
        filepath = os.path.join(storage_dir, filename)
        doc.save(filepath)

        return filepath


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text with simple bold (**text**) and italic (*text*) formatting."""
    import re

    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

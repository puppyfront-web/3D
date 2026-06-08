"""Export service for Word, PDF and PPTX generation."""

import os
from abc import ABC, abstractmethod
from datetime import datetime, timezone
from typing import Optional

from app.core.config import settings


class ExportService(ABC):
    """Abstract export service."""

    @abstractmethod
    async def export_to_word(
        self,
        content: str,
        filename: str,
        title: Optional[str] = None,
    ) -> str:
        """Export content to a Word document. Returns the file path."""

    @abstractmethod
    async def export_to_pdf(
        self,
        content: str,
        filename: str,
        title: Optional[str] = None,
    ) -> str:
        """Export content to a PDF document. Returns the file path."""

    @abstractmethod
    async def export_to_pptx(
        self,
        content: str,
        filename: str,
        title: Optional[str] = None,
    ) -> str:
        """Export content to a PowerPoint presentation. Returns the file path."""


class LocalExportService(ExportService):
    """Export service that writes files to local storage."""

    async def export_to_word(
        self,
        content: str,
        filename: str,
        title: Optional[str] = None,
    ) -> str:
        """Export markdown-like content to a Word document."""
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

        storage_dir = os.path.abspath(settings.storage_path)
        os.makedirs(storage_dir, exist_ok=True)
        filepath = os.path.join(storage_dir, filename)
        doc.save(filepath)
        return filepath

    async def export_to_pdf(
        self,
        content: str,
        filename: str,
        title: Optional[str] = None,
    ) -> str:
        """Export markdown-like content to a PDF document."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from reportlab.lib.colors import HexColor

        storage_dir = os.path.abspath(settings.storage_path)
        os.makedirs(storage_dir, exist_ok=True)
        filepath = os.path.join(storage_dir, filename)

        doc = SimpleDocTemplate(filepath, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 0.2 * inch))
            elif stripped.startswith("# "):
                story.append(Paragraph(stripped[2:], styles["Heading1"]))
            elif stripped.startswith("## "):
                story.append(Paragraph(stripped[3:], styles["Heading2"]))
            elif stripped.startswith("### "):
                story.append(Paragraph(stripped[4:], styles["Heading3"]))
            elif stripped.startswith("- ") or stripped.startswith("* "):
                safe = stripped[2:].replace("&", "&amp;").replace("<", "&lt;")
                story.append(Paragraph(f"&bull; {safe}", styles["Normal"]))
            else:
                safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, styles["Normal"]))

        doc.build(story)
        return filepath

    async def export_to_pptx(
        self,
        content: str,
        filename: str,
        title: Optional[str] = None,
    ) -> str:
        """Export markdown-like content to a PowerPoint presentation."""
        from app.exporters.pptx_exporter import PPTXExporter

        exporter = PPTXExporter()
        return await exporter.export(
            content=content,
            filename=filename,
            title=title,
        )


def get_export_service() -> ExportService:
    """Factory function to create the export service."""
    return LocalExportService()

"""Exports router — export proposals to Word/PDF/PPTX."""

import os
import uuid
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.exceptions import NotFoundException
from app.db.session import get_db
from app.models.generation import GenerationOutput, GenerationTask
from app.schemas.common import Response

router = APIRouter(prefix="/exports", tags=["exports"])


async def _get_task_output(generation_id: uuid.UUID, db: AsyncSession):
    """Look up a GenerationTask and output.

    The canonical ID is GenerationTask.id. For frontend compatibility, also
    accept GenerationOutput.id and resolve it back to its parent task.
    """
    task = await db.get(GenerationTask, generation_id)
    output = None

    if task:
        result = await db.execute(
            select(GenerationOutput).where(GenerationOutput.task_id == generation_id)
        )
        output = result.scalars().first()
    else:
        output = await db.get(GenerationOutput, generation_id)
        if output:
            task = await db.get(GenerationTask, output.task_id)

    if not task:
        raise NotFoundException("GenerationTask or GenerationOutput", str(generation_id))
    if not output:
        raise NotFoundException("GenerationOutput for task", str(task.id))
    return task, output


def _check_export_eligibility(output: GenerationOutput) -> tuple[bool, list[str]]:
    """Check if a proposal is eligible for export (all sections approved).

    Returns (is_eligible, list_of_blocker_messages).
    """
    blockers: list[str] = []
    meta = output.sections_meta

    if not meta:
        # No sections_meta means proposal was generated before HITL was added — allow export
        return True, []

    for section in meta:
        if section.get("status") != "approved":
            blockers.append(f"章节「{section.get('title', '?')}」未审核通过")
        if section.get("require_human_review") and not section.get("human_confirmed"):
            blockers.append(f"章节「{section.get('title', '?')}」需人工确认")

    return len(blockers) == 0, blockers


@router.post("/word/{task_id}")
async def export_to_word(
    task_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """Export a generation task's output as a Word document."""
    task, output = await _get_task_output(task_id, db)

    eligible, blockers = _check_export_eligibility(output)
    if not eligible:
        raise HTTPException(
            status_code=403,
            detail={"message": "导出前需完成审核", "blockers": blockers},
        )

    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")

    doc = DocxDocument()

    # Style setup
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Parse content — handle markdown-like text
    content = output.content
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        elif stripped.startswith("# "):
            heading = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            heading = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            heading = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("---"):
            doc.add_paragraph("-" * 50)
        elif stripped.startswith("| "):
            doc.add_paragraph(stripped, style="Normal")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith(tuple(f"{i}. " for i in range(1, 10))):
            doc.add_paragraph(stripped, style="List Number")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        else:
            doc.add_paragraph(stripped)

    # Add footer
    doc.add_paragraph("")
    footer_para = doc.add_paragraph(
        f"Generated on {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} by 3D Wall Platform"
    )
    footer_para.style = doc.styles["Normal"]
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    storage_dir = os.path.abspath(settings.storage_path)
    os.makedirs(storage_dir, exist_ok=True)
    filename = f"export_{task_id}.docx"
    filepath = os.path.join(storage_dir, filename)
    doc.save(filepath)

    return FileResponse(
        path=filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"proposal_{task.type}.docx",
    )


@router.post("/pdf/{task_id}")
async def export_to_pdf(
    task_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """Export a generation task's output as a PDF document."""
    task, output = await _get_task_output(task_id, db)

    eligible, blockers = _check_export_eligibility(output)
    if not eligible:
        raise HTTPException(
            status_code=403,
            detail={"message": "导出前需完成审核", "blockers": blockers},
        )

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from reportlab.lib.colors import HexColor
    except ImportError:
        raise HTTPException(status_code=500, detail="reportlab not installed")

    storage_dir = os.path.abspath(settings.storage_path)
    os.makedirs(storage_dir, exist_ok=True)
    filename = f"export_{task_id}.pdf"
    filepath = os.path.join(storage_dir, filename)

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CustomH1",
            parent=styles["Heading1"],
            fontSize=20,
            textColor=HexColor("#1a73e8"),
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CustomH2",
            parent=styles["Heading2"],
            fontSize=16,
            textColor=HexColor("#333333"),
            spaceAfter=8,
        )
    )

    story = []
    for line in output.content.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.2 * inch))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], styles["CustomH1"]))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], styles["CustomH2"]))
        elif stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], styles["Heading3"]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            story.append(Paragraph(f"&bull; {stripped[2:]}", styles["Normal"]))
        elif stripped.startswith("---"):
            story.append(Spacer(1, 0.1 * inch))
        else:
            safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, styles["Normal"]))

    doc.build(story)

    return FileResponse(
        path=filepath,
        media_type="application/pdf",
        filename=f"proposal_{task.type}.pdf",
    )


@router.post("/pptx/{task_id}")
async def export_to_pptx(
    task_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
):
    """Export a generation task's output as a PowerPoint presentation."""
    task, output = await _get_task_output(task_id, db)

    eligible, blockers = _check_export_eligibility(output)
    if not eligible:
        raise HTTPException(
            status_code=403,
            detail={"message": "导出前需完成审核", "blockers": blockers},
        )

    try:
        from app.exporters.pptx_exporter import PPTXExporter
    except ImportError:
        raise HTTPException(status_code=500, detail="python-pptx not installed")

    exporter = PPTXExporter()
    filename = f"export_{task_id}.pptx"
    filepath = await exporter.export(
        content=output.content,
        filename=filename,
        title=task.type or "方案报告",
    )

    return FileResponse(
        path=filepath,
        media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        filename=f"proposal_{task.type}.pptx",
    )
